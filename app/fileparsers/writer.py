"""Format-aware document writing for FileSystemTool.

Single public entry point: :func:`write_document`.

Dispatch is by file extension:

- Plain text (``.txt``, ``.md``, ``.markdown``, ``.html``, ``.htm``, ``.csv``)
  -- UTF-8 text, preserving the previous ``Path.write_text`` behavior.
- Word (``.docx``) -- python-docx, lightweight markdown -> docx conversion
  (headings, bold, italic, bullet/numbered lists, inline code, code blocks,
  horizontal rules, and simple markdown tables).
- Excel (``.xlsx``) -- openpyxl; markdown tables become worksheets, otherwise
  rows split on ``|`` / comma / tab in a single worksheet.
- PDF (``.pdf``) -- PyMuPDF (fitz); readable text with heading sizes and
  automatic pagination. No images, no advanced layout.
- Unknown extensions fall back to UTF-8 text (previous behavior).

No new dependencies are introduced.
"""

from __future__ import annotations

import logging
import os
import re
from pathlib import Path

logger = logging.getLogger(__name__)

PLAIN_TEXT_EXTENSIONS = frozenset({".txt", ".md", ".markdown", ".html", ".htm"})

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_RULE_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_SEPARATOR_CELL_RE = re.compile(r"^:?-{1,}:?$")
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def write_document(path: Path, content: str) -> tuple[str, int]:
    """Write ``content`` to ``path`` in the format implied by its extension.

    Returns ``(format_name, written_bytes)``. Non-string content is converted
    to ``str`` before writing.
    """
    path = Path(path)
    if not isinstance(content, str):
        content = str(content)

    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _write_docx(path, content)
    if suffix == ".xlsx":
        return _write_xlsx(path, content)
    if suffix == ".pdf":
        return _write_pdf(path, content)
    if suffix == ".csv":
        return _write_csv(path, content)
    return _write_text(path, content)


# ---------------------------------------------------------------------------
# CSV (.csv)
# ---------------------------------------------------------------------------


def _write_csv(path: Path, content: str) -> tuple[str, int]:
    import csv

    tables = [block["rows"] for block in _iter_blocks(content) if block["kind"] == "table"]

    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if tables:
            for idx, table in enumerate(tables):
                if idx > 0:
                    writer.writerow([])
                for row in table:
                    writer.writerow(row)
        else:
            for line in content.splitlines():
                if line.strip():
                    writer.writerow(_split_row(line))

    return "csv", path.stat().st_size


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


def _write_text(path: Path, content: str) -> tuple[str, int]:
    path.write_text(content, encoding="utf-8")
    return "text", path.stat().st_size


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------


def _write_docx(path: Path, content: str) -> tuple[str, int]:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    for block in _iter_blocks(content):
        kind = block["kind"]
        if kind == "heading":
            document.add_heading(block["text"], level=min(block["level"], 9))
        elif kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "bullet":
            _add_rich_paragraph(document.add_paragraph(style="List Bullet"), block["text"])
        elif kind == "numbered":
            _add_rich_paragraph(document.add_paragraph(style="List Number"), block["text"])
        elif kind == "rule":
            _add_horizontal_rule(document)
        elif kind == "table":
            _add_table(document, block["rows"])
        else:
            _add_rich_paragraph(document.add_paragraph(), block["text"])

    document.save(str(path))
    return "docx", path.stat().st_size


def _add_rich_paragraph(paragraph, text: str) -> None:
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_horizontal_rule(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(p_bdr)


def _add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell


# ---------------------------------------------------------------------------
# Excel (.xlsx)
# ---------------------------------------------------------------------------


def _write_xlsx(path: Path, content: str) -> tuple[str, int]:
    from openpyxl import Workbook

    tables = [block["rows"] for block in _iter_blocks(content) if block["kind"] == "table"]

    workbook = Workbook()
    if tables:
        workbook.remove(workbook.active)
        for index, table in enumerate(tables):
            worksheet = workbook.create_sheet(title=f"Sheet{index + 1}")
            for row in table:
                worksheet.append(row)
    else:
        worksheet = workbook.active
        worksheet.title = "Sheet1"
        for line in content.splitlines():
            if line.strip():
                worksheet.append(_split_row(line))

    workbook.save(str(path))
    return "xlsx", path.stat().st_size


def _split_row(line: str) -> list[str]:
    line = line.strip()
    if "|" in line:
        return [cell.strip() for cell in line.strip("|").split("|")]
    if "," in line:
        return [cell.strip() for cell in line.split(",")]
    if "\t" in line:
        return [cell.strip() for cell in line.split("\t")]
    return [line]


# ---------------------------------------------------------------------------
# PDF (.pdf)
# ---------------------------------------------------------------------------

_PDF_FONT_CANDIDATES = (
    "C:/Windows/Fonts/msyh.ttc",
    "C:/Windows/Fonts/simsun.ttc",
    "C:/Windows/Fonts/segoeui.ttf",
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/arialuni.ttf",
)


def _write_pdf(path: Path, content: str) -> tuple[str, int]:
    import fitz

    fontfile, fontname, font = _pdf_font_for(content)
    document = fitz.open()
    page = document.new_page()
    margin = 72
    y = margin
    page_bottom = page.rect.height - margin

    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            text, size = stripped[4:].strip(), 13
        elif stripped.startswith("## "):
            text, size = stripped[3:].strip(), 16
        elif stripped.startswith("# "):
            text, size = stripped[2:].strip(), 20
        elif stripped.startswith("```"):
            continue
        elif stripped in ("---", "***", "___"):
            text, size = "-" * 48, 11
        else:
            text, size = line, 11

        for display_line in _wrap_text(text, size, page.rect.width - 2 * margin, font=font, fontname=fontname):
            if y + size > page_bottom:
                page = document.new_page()
                y = margin
            if fontfile:
                page.insert_text(
                    (margin, y), display_line, fontsize=size, fontfile=fontfile, fontname=fontname
                )
            else:
                page.insert_text((margin, y), display_line, fontsize=size, fontname=fontname)
            y += size * 1.5

    document.save(str(path))
    document.close()
    return "pdf", path.stat().st_size


def _pdf_font_for(content: str) -> tuple[str | None, str, "object | None"]:
    """Select a font for the PDF writer.

    Pure-ASCII content uses the built-in Helvetica (no embedding). Non-ASCII
    content embeds the first system font covering the required characters, so
    UTF-8 text (accents, CJK, etc.) renders correctly. Non-BMP symbols (emoji)
    are limited by available fonts. ``SAMAKTHA_PDF_FONT`` overrides the search.
    """
    import fitz

    non_ascii = sorted({c for c in content if ord(c) > 127})
    if not non_ascii:
        return None, "helv", None

    candidates = list(_PDF_FONT_CANDIDATES)
    override = os.environ.get("SAMAKTHA_PDF_FONT")
    if override:
        candidates.insert(0, override)

    fallback = None
    for candidate in candidates:
        if not candidate or not os.path.exists(candidate):
            continue
        try:
            font = fitz.Font(fontfile=candidate)
        except Exception:
            continue
        if all(font.has_glyph(ord(char)) for char in non_ascii):
            return candidate, "embed", font
        if fallback is None:
            fallback = (candidate, font)

    if fallback:
        return fallback[0], "embed", fallback[1]
    return None, "helv", None


def _wrap_text(text: str, size: float, max_width: float, font=None, fontname: str = "helv") -> list[str]:
    import fitz

    if not text:
        return [""]

    def measure(value: str) -> float:
        if font is not None:
            return font.text_length(value, fontsize=size)
        return fitz.get_text_length(value, fontname=fontname, fontsize=size)

    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        if measure(trial) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


# ---------------------------------------------------------------------------
# Lightweight markdown block parser (shared by docx/xlsx writers)
# ---------------------------------------------------------------------------


def _iter_blocks(content: str):
    """Yield markdown blocks as dicts with a ``kind`` key.

    Kinds: heading, paragraph, bullet, numbered, code, rule, table.
    Unknown markdown collapses into paragraph blocks.
    """
    lines = content.splitlines()
    index = 0
    count = len(lines)

    while index < count:
        raw = lines[index]
        stripped = raw.strip()

        # Fenced code block
        if stripped.startswith("```"):
            buffer: list[str] = []
            index += 1
            while index < count and not lines[index].strip().startswith("```"):
                buffer.append(lines[index])
                index += 1
            yield {"kind": "code", "text": "\n".join(buffer)}
            index += 1
            continue

        # Heading (levels 1-3 become headings; deeper levels fall back to text)
        heading = _HEADING_RE.match(stripped)
        if heading and len(heading.group(1)) <= 3:
            yield {"kind": "heading", "level": len(heading.group(1)), "text": heading.group(2).strip()}
            index += 1
            continue

        # Horizontal rule
        if _RULE_RE.match(stripped):
            yield {"kind": "rule"}
            index += 1
            continue

        # Markdown table (current line is a header, next line is a separator)
        if "|" in raw and index + 1 < count and _is_separator_row(lines[index + 1]):
            table_lines = [raw]
            index += 2
            while index < count and "|" in lines[index]:
                table_lines.append(lines[index])
                index += 1
            rows = [
                [cell.strip() for cell in line.strip().strip("|").split("|")]
                for line in table_lines
            ]
            yield {"kind": "table", "rows": rows}
            continue

        # Bullet list
        bullet = _BULLET_RE.match(stripped)
        if bullet:
            yield {"kind": "bullet", "text": bullet.group(1).strip()}
            index += 1
            continue

        # Numbered list
        numbered = _NUMBERED_RE.match(stripped)
        if numbered:
            yield {"kind": "numbered", "text": numbered.group(1).strip()}
            index += 1
            continue

        # Blank line separates paragraphs
        if not stripped:
            index += 1
            continue

        # Paragraph: accumulate consecutive non-block lines
        buffer = [raw]
        index += 1
        while index < count and not _starts_block(lines[index]):
            buffer.append(lines[index])
            index += 1
        yield {"kind": "paragraph", "text": "\n".join(part for part in buffer)}


def _is_separator_row(line: str) -> bool:
    if "|" not in line:
        return False
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(_SEPARATOR_CELL_RE.match(cell) for cell in cells)


def _starts_block(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith("```"):
        return True
    if _HEADING_RE.match(stripped) or _RULE_RE.match(stripped):
        return True
    if _BULLET_RE.match(stripped) or _NUMBERED_RE.match(stripped):
        return True
    if "|" in stripped:
        return True
    return False

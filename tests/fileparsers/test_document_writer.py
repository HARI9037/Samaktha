"""Tests for app.fileparsers.writer.write_document.

Covers plain text, markdown, docx, xlsx, pdf, unknown-extension fallback,
empty files, unicode, large content, absolute paths, and round-trip
validation through the existing document parser chain.
"""

import pytest

from app.fileparsers.factory import DocumentParserFactory
from app.fileparsers.writer import write_document

SAMPLE_MARKDOWN = """# Project Report

This is a **bold** statement with *italic* and `inline code`.

## Details

- bullet one
- bullet two

1. first
2. second

---

| Name | Age |
|------|-----|
| Alice | 30 |
| Bob | 25 |

```python
print("hello")
```
"""


def test_write_txt_returns_format_and_bytes(tmp_path):
    target = tmp_path / "plain.txt"
    fmt, written = write_document(target, "hello, I'm Samaktha")
    assert fmt == "text"
    assert written == target.stat().st_size
    assert target.read_text(encoding="utf-8") == "hello, I'm Samaktha"


def test_write_markdown_round_trips_through_parser(tmp_path):
    target = tmp_path / "notes.md"
    fmt, _ = write_document(target, SAMPLE_MARKDOWN)
    assert fmt == "text"
    result = DocumentParserFactory.parse(target)
    assert result.ok
    assert "# Project Report" in result.text
    assert "bullet two" in result.text


def test_write_docx_is_valid_and_renders_markdown(tmp_path):
    from docx import Document

    target = tmp_path / "report.docx"
    fmt, written = write_document(target, SAMPLE_MARKDOWN)
    assert fmt == "docx"
    assert written == target.stat().st_size

    document = Document(str(target))
    paragraph_texts = [p.text for p in document.paragraphs]
    assert "Project Report" in paragraph_texts
    assert any("bold" in text for text in paragraph_texts)
    assert any(text.startswith("bullet") for text in paragraph_texts)
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Name"
    assert document.tables[0].cell(1, 1).text == "30"


def test_write_docx_applies_bold_run(tmp_path):
    from docx import Document

    target = tmp_path / "bold.docx"
    write_document(target, "World **bold** text")
    document = Document(str(target))
    runs = document.paragraphs[0].runs
    assert any(run.bold and run.text == "bold" for run in runs)


def test_write_xlsx_markdown_tables_become_worksheets(tmp_path):
    from openpyxl import load_workbook

    target = tmp_path / "data.xlsx"
    fmt, _ = write_document(target, SAMPLE_MARKDOWN)
    assert fmt == "xlsx"

    workbook = load_workbook(str(target))
    assert "Sheet1" in workbook.sheetnames
    worksheet = workbook["Sheet1"]
    rows = [list(row) for row in worksheet.iter_rows(values_only=True)]
    assert rows[0] == ["Name", "Age"]
    assert ["Alice", "30"] in rows


def test_write_xlsx_non_table_splits_rows(tmp_path):
    from openpyxl import load_workbook

    target = tmp_path / "list.xlsx"
    write_document(target, "name,email\nAlice,alice@example.com\nBob,bob@example.com")
    workbook = load_workbook(str(target))
    worksheet = workbook.active
    rows = [list(row) for row in worksheet.iter_rows(values_only=True)]
    assert rows[0] == ["name", "email"]
    assert rows[1] == ["Alice", "alice@example.com"]


def test_write_xlsx_single_column_when_no_separators(tmp_path):
    from openpyxl import load_workbook

    target = tmp_path / "single.xlsx"
    write_document(target, "alpha\nbeta")
    workbook = load_workbook(str(target))
    rows = [list(row) for row in workbook.active.iter_rows(values_only=True)]
    assert rows == [["alpha"], ["beta"]]


def test_write_pdf_is_readable_text(tmp_path):
    import fitz

    target = tmp_path / "doc.pdf"
    fmt, _ = write_document(target, SAMPLE_MARKDOWN)
    assert fmt == "pdf"

    document = fitz.open(str(target))
    assert document.page_count >= 1
    text = "".join(page.get_text() for page in document)
    assert "Project Report" in text
    assert "bullet one" in text
    document.close()


def test_write_unknown_extension_falls_back_to_text(tmp_path):
    target = tmp_path / "archive.dat"
    fmt, _ = write_document(target, "raw bytes as text")
    assert fmt == "text"
    assert target.read_text(encoding="utf-8") == "raw bytes as text"


def test_write_empty_files(tmp_path):
    text_target = tmp_path / "empty.txt"
    fmt, written = write_document(text_target, "")
    assert fmt == "text"
    assert written == 0

    from docx import Document
    docx_target = tmp_path / "empty.docx"
    write_document(docx_target, "")
    assert Document(str(docx_target)).paragraphs == []

    from openpyxl import load_workbook
    xlsx_target = tmp_path / "empty.xlsx"
    write_document(xlsx_target, "")
    assert load_workbook(str(xlsx_target)).active.max_row == 1

    import fitz
    pdf_target = tmp_path / "empty.pdf"
    write_document(pdf_target, "")
    document = fitz.open(str(pdf_target))
    assert document.page_count == 1
    document.close()


def test_write_unicode_text(tmp_path):
    content = "héllo wörld — 日本語"
    text_target = tmp_path / "unicode.txt"
    write_document(text_target, content)
    assert text_target.read_text(encoding="utf-8") == content

    from docx import Document
    docx_target = tmp_path / "unicode.docx"
    write_document(docx_target, content)
    assert "日本語" in "".join(p.text for p in Document(str(docx_target)).paragraphs)

    import fitz
    pdf_target = tmp_path / "unicode.pdf"
    write_document(pdf_target, content)
    document = fitz.open(str(pdf_target))
    assert "日本語" in "".join(page.get_text() for page in document)
    document.close()


def test_write_large_files(tmp_path):
    content = "\n".join(f"line {i}: some content to pad out the document" for i in range(5000))

    text_target = tmp_path / "large.txt"
    write_document(text_target, content)
    assert "line 0:" in text_target.read_text(encoding="utf-8")
    assert "line 4999:" in text_target.read_text(encoding="utf-8")

    from docx import Document
    docx_target = tmp_path / "large.docx"
    write_document(docx_target, content)
    assert "line 4999:" in "".join(p.text for p in Document(str(docx_target)).paragraphs)

    import fitz
    pdf_target = tmp_path / "large.pdf"
    write_document(pdf_target, content)
    document = fitz.open(str(pdf_target))
    assert document.page_count > 1
    document.close()


def test_write_absolute_path_writes_exactly_there(tmp_path):
    target = tmp_path / "sub" / "abs.txt"
    target.parent.mkdir(parents=True)
    write_document(target, "content")
    assert target.exists()
    assert target.read_text(encoding="utf-8") == "content"


def test_write_non_string_content_is_converted(tmp_path):
    target = tmp_path / "number.txt"
    fmt, _ = write_document(target, 12345)
    assert fmt == "text"
    assert target.read_text(encoding="utf-8") == "12345"


def test_round_trip_through_existing_parser(tmp_path):
    """Round-trip via the existing parser chain for torch-free formats.

    docx/xlsx read-back depends on docling/torch, whose DLL initialization is
    flaky under a full-suite process on Windows; those formats are validated
    directly with the producing libraries instead (see docx/xlsx tests above).
    """
    payload = "# Round Trip\n\n| A | B |\n|---|---|\n| 1 | 2 |\n"
    for extension in ("txt", "md", "pdf"):
        target = tmp_path / f"rt.{extension}"
        write_document(target, payload)
        result = DocumentParserFactory.parse(target)
        assert result.ok, f"{extension}: {result.error}"
        assert "Round Trip" in result.text, f"{extension} missing content"

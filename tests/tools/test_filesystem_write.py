"""Tests for FileSystemTool write-path behavior and format dispatch.

Covers:
- absolute write paths preserved exactly
- relative write paths resolve against the configured default output dir
  (never the repository root / cwd)
- format dispatch by extension (txt/md/docx/xlsx/pdf) through the tool
- existing overwrite semantics
"""

import pytest

from app.tools.base import ToolResult
from app.tools.filesystem import FileSystemTool

pytestmark = pytest.mark.asyncio


@pytest.fixture
def write_dir(tmp_path):
    return tmp_path / "output"


@pytest.fixture
def tool(write_dir):
    return FileSystemTool(write_dir=str(write_dir))


async def test_write_absolute_path_preserved_exactly(tool, tmp_path):
    target = tmp_path / "notes" / "abs.txt"
    result = await tool.run({"action": "write", "path": str(target), "content": "hello"})
    assert result.ok, result.error
    assert result.data["path"] == str(target)
    assert target.read_text(encoding="utf-8") == "hello"


async def test_write_relative_path_uses_default_output_dir(tool, write_dir):
    result = await tool.run({"action": "write", "path": "hello.txt", "content": "hi"})
    assert result.ok, result.error
    written = write_dir / "hello.txt"
    assert result.data["path"] == str(written)
    assert written.exists()


async def test_write_relative_path_with_subdir_uses_default_output_dir(tool, write_dir):
    result = await tool.run({"action": "write", "path": "reports/notes.md", "content": "# n"})
    assert result.ok, result.error
    written = write_dir / "reports" / "notes.md"
    assert result.data["path"] == str(written)
    assert written.exists()


async def test_write_relative_path_never_anchors_to_cwd(tmp_path, monkeypatch):
    import app.memory.resources as resources_module
    import app.tools.resolver as resolver_module

    class _StubRegistry:
        def __init__(self):
            self.registered = []

        def register(self, path):
            self.registered.append(str(path))

        def lookup(self, path):
            return None

        def remove(self, path):
            return None

    monkeypatch.setattr(resources_module, "ResourceRegistry", _StubRegistry)
    monkeypatch.setattr(resolver_module, "ResourceRegistry", _StubRegistry)

    cwd_root = tmp_path / "cwd"
    cwd_root.mkdir()
    write_dir = tmp_path / "output"
    tool = FileSystemTool(write_dir=str(write_dir))
    monkeypatch.chdir(cwd_root)

    result = await tool.run({"action": "write", "path": "file.txt", "content": "x"})
    assert result.ok, result.error
    assert not (cwd_root / "file.txt").exists()
    assert (write_dir / "file.txt").exists()


async def test_write_returns_format_and_written_bytes(tool, write_dir):
    result = await tool.run({"action": "write", "path": "data.md", "content": "# Hello"})
    assert result.data["format"] == "text"
    assert result.data["written_bytes"] == len("# Hello".encode("utf-8"))


async def test_write_docx_via_tool_round_trips(tool, write_dir):
    from docx import Document

    result = await tool.run(
        {"action": "write", "path": "report.docx", "content": "# Title\n\nBody text."}
    )
    assert result.ok, result.error
    assert result.data["format"] == "docx"
    written = write_dir / "report.docx"
    assert written.exists()

    document = Document(str(written))
    assert any("Title" in p.text for p in document.paragraphs)
    assert any("Body text" in p.text for p in document.paragraphs)


async def test_write_xlsx_via_tool_round_trips(tool, write_dir):
    from openpyxl import load_workbook

    result = await tool.run(
        {"action": "write", "path": "data.xlsx", "content": "name,age\nAlice,30\n"}
    )
    assert result.ok, result.error
    assert result.data["format"] == "xlsx"
    written = write_dir / "data.xlsx"
    assert written.exists()

    workbook = load_workbook(str(written))
    rows = [list(row) for row in workbook.active.iter_rows(values_only=True)]
    assert rows == [["name", "age"], ["Alice", "30"]]


async def test_write_pdf_via_tool_round_trips(tool, write_dir):
    result = await tool.run(
        {"action": "write", "path": "doc.pdf", "content": "# Report\n\nSome body."}
    )
    assert result.ok, result.error
    assert result.data["format"] == "pdf"

    read_back = await tool.run({"action": "read", "path": str(write_dir / "doc.pdf")})
    assert read_back.ok, read_back.error
    assert "Report" in read_back.data.get("content", "")


async def test_write_overwrites_existing_file(tool, write_dir):
    first = await tool.run({"action": "write", "path": "same.txt", "content": "v1"})
    assert first.ok
    second = await tool.run({"action": "write", "path": "same.txt", "content": "v2-longer"})
    assert second.ok
    assert (write_dir / "same.txt").read_text(encoding="utf-8") == "v2-longer"


async def test_write_converts_non_string_content(tool, write_dir):
    result = await tool.run({"action": "write", "path": "num.txt", "content": 42})
    assert result.ok, result.error
    assert (write_dir / "num.txt").read_text(encoding="utf-8") == "42"


async def test_write_unknown_extension_falls_back_to_text(tool, write_dir):
    result = await tool.run({"action": "write", "path": "blob.dat", "content": "text"})
    assert result.ok, result.error
    assert result.data["format"] == "text"
    assert (write_dir / "blob.dat").read_text(encoding="utf-8") == "text"
